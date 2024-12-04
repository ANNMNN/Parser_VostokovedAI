import re
from docx import Document
from bs4 import BeautifulSoup
from PyPDF2 import PdfReader
import fitz  # PyMuPDF -djvu
from parser.table_processor import process_table
from parser.image_processor import process_image
from PIL import Image
from io import BytesIO

# очистка данных
def clean_parsed_data(data):

    def clean_text(text, to_lower=True):
        text = re.sub(r"\s+", " ", text.strip())
        text = re.sub(r"\.{2,}", ".", text)
        text = re.sub(r"[^\w\s.,!?-]", "", text)
        if to_lower:
            text = text.lower()
        return text

    def clean_table(table):
        cleaned_table = []
        seen_rows = set()
        for r in table:
            cleaned_row = [clean_text(cell) for cell in r if cell.strip()]
            if cleaned_row and tuple(cleaned_row) not in seen_rows:
                cleaned_table.append(cleaned_row)
                seen_rows.add(tuple(cleaned_row))
        return cleaned_table

    def filter_tables(tables):
        return [table for table in tables if table]

    cleaned_data = data.copy()
    cleaned_data["text"] = [clean_text(text) for text in cleaned_data["text"] if clean_text(text)]
    cleaned_data["tables"] = filter_tables([clean_table(table) for table in cleaned_data["tables"]])
    cleaned_data["images"] = [process_image(img) for img in cleaned_data.get("images", []) if img]
    return cleaned_data


class FileParser:
    def __init__(self, file_path):
        self.file_path = file_path

    # проверка форматов
    def parse(self):
        if self.file_path.endswith(".docx") or self.file_path.endswith(".doc"):
            data = self.parse_docx()
        elif self.file_path.endswith(".html") or self.file_path.endswith(".htm"):
            data = self.parse_html()
        elif self.file_path.endswith(".pdf"):
            data = self.parse_pdf()
        elif self.file_path.endswith(".djvu"):
            data = self.parse_djvu()
        else:
            raise ValueError("Неподдерживаемый формат файла")
        return self.split_into_chunks(data)

    # парсинг docx
    def parse_docx(self):
        document = Document(self.file_path)
        data = {"text": [], "tables": [], "images": []}

        for p in document.paragraphs:
            if p.text.strip():
                data["text"].append(p.text)

        for t in document.tables:
            table_data = []
            for row in t.rows:
                table_data.append([cell.text.strip() for cell in row.cells])
            if table_data:
                data["tables"].append(table_data)



        data["images"] = self.extract_images_from_docx(document)

        return clean_parsed_data(data)

    # изоображение из docx
    def extract_images_from_docx(self, document):
        images = []
        for rel in document.part.rels.values():
            if "image" in rel.target_ref:  # Проверка на изображение
                try:
                    image_data = rel.target_part.blob
                    image = Image.open(BytesIO(image_data))
                    images.append(image)
                except Exception as e:
                    print(f"Ошибка при обработке изображения: {e}")
        return images

    # парсинг pdf
    def parse_pdf(self):
        reader = PdfReader(self.file_path)
        text = []
        for p in reader.pages:
            text.append(p.extract_text())
        return clean_parsed_data({"text": text, "tables": [], "images": []})

    # парсинг djvu
    def parse_djvu(self):
        doc = fitz.open(self.file_path)
        text = []
        imgs = []
        for p_num in range(len(doc)):
            page = doc.load_page(p_num)
            text.append(page.get_text())
            for i in enumerate(page.get_images(full=True)):
                xref = i[0]
                base_img = doc.extract_image(xref)
                img_data = base_img["image"]
                img = Image.open(BytesIO(img_data))
                imgs.append(img)
        return clean_parsed_data({"text": text, "tables": [], "images": imgs})

    # деление на чанки
    def split_into_chunks(self, data, chunk_size=500):
        chunks = []
        current_chunk = {"text": "", "tables": [], "images": []}
        current_length = 0

        # Добавление текста в чанки
        for tx in data["text"]:
            for sentence in re.split(r"(?<=[.!?])\s+", tx):
                if current_length + len(sentence) > chunk_size:
                    chunks.append(current_chunk)
                    current_chunk = {"text": "", "tables": [], "images": []}
                    current_length = 0
                current_chunk["text"] += sentence + " "
                current_length += len(sentence)

        # Добавляем текущий текстовый чанк, если он непустой
        if current_chunk["text"].strip() or current_chunk["tables"] or current_chunk["images"]:
            chunks.append(current_chunk)
            current_chunk = {"text": "", "tables": [], "images": []}
            current_length = 0

        # Добавление таблиц в чанки
        for table in data["tables"]:
            if current_length + len(str(table)) > chunk_size:
                chunks.append(current_chunk)
                current_chunk = {"text": "", "tables": [], "images": []}
                current_length = 0
            current_chunk["tables"].append(table)
            current_length += len(str(table))

        # Добавляем текущий чанк, если он непустой
        if current_chunk["text"].strip() or current_chunk["tables"] or current_chunk["images"]:
            chunks.append(current_chunk)
            current_chunk = {"text": "", "tables": [], "images": []}
            current_length = 0

        # Добавление изображений в чанки
        for image in data["images"]:
            if current_length + 1 > chunk_size:  # 1 условно для размера изображения
                chunks.append(current_chunk)
                current_chunk = {"text": "", "tables": [], "images": []}
                current_length = 0
            current_chunk["images"].append(image)
            current_length += 1

        # Добавляем последний чанк, если он непустой
        if current_chunk["text"].strip() or current_chunk["tables"] or current_chunk["images"]:
            chunks.append(current_chunk)

        return chunks

